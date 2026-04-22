"""Report rendering helpers for workflow presentation."""

from __future__ import annotations

import json
import textwrap
from pathlib import Path
from typing import Any, Sequence

import matplotlib.pyplot as plt
from matplotlib.backends.backend_pdf import PdfPages


def normalize_report_formats(value: str | None) -> tuple[str, ...]:
    """Return normalized report formats."""
    mode = str(value or "both").strip().lower()
    if mode == "pdf":
        return ("pdf",)
    if mode == "docx":
        return ("docx",)
    return ("pdf", "docx")


def _normalize_payload(payload: dict[str, Any]) -> dict[str, Any]:
    out = dict(payload)
    out["title"] = str(out.get("title") or "Analysis Report")
    out["subtitle"] = str(out.get("subtitle") or "")
    sections = out.get("sections")
    out["sections"] = sections if isinstance(sections, list) else []
    figures = out.get("figures")
    out["figures"] = figures if isinstance(figures, list) else []
    return out


def _to_json_text(value: Any) -> str:
    return json.dumps(value, indent=2, sort_keys=True)


def _stringify_cell(value: Any) -> str:
    if isinstance(value, (dict, list, tuple)):
        try:
            return json.dumps(value, sort_keys=True)
        except Exception:
            return str(value)
    return str(value)


def _section_table_data(section: dict[str, Any]) -> tuple[list[str], list[list[str]]] | None:
    table = section.get("table")
    if isinstance(table, dict):
        headers_raw = table.get("headers")
        rows_raw = table.get("rows")
        if isinstance(headers_raw, list) and isinstance(rows_raw, list):
            headers = [str(h) for h in headers_raw]
            rows: list[list[str]] = []
            for row in rows_raw:
                if isinstance(row, dict):
                    rows.append([_stringify_cell(row.get(h, "")) for h in headers])
                elif isinstance(row, (list, tuple)):
                    vals = list(row)
                    if len(vals) < len(headers):
                        vals.extend([""] * (len(headers) - len(vals)))
                    rows.append([_stringify_cell(v) for v in vals[: len(headers)]])
                else:
                    rows.append([_stringify_cell(row)] + [""] * max(0, len(headers) - 1))
            return headers, rows

    key_values = section.get("key_values")
    if isinstance(key_values, dict) and key_values:
        headers = ["Metric", "Value"]
        rows = [[str(k), _stringify_cell(v)] for k, v in key_values.items()]
        return headers, rows

    return None


def _table_col_widths(n_cols: int) -> list[float]:
    if n_cols <= 1:
        return [1.0]
    if n_cols == 2:
        return [0.30, 0.70]
    if n_cols == 3:
        return [0.20, 0.16, 0.64]
    lead = [0.15] * (n_cols - 1)
    last = max(0.25, 1.0 - sum(lead))
    return lead + [last]


def _wrap_table_rows(rows: list[list[str]], col_widths: list[float]) -> list[list[str]]:
    wrapped: list[list[str]] = []
    if not rows:
        return wrapped
    total_chars = 96
    for row in rows:
        out_row: list[str] = []
        for i, value in enumerate(row):
            width = int(max(8, (col_widths[i] if i < len(col_widths) else col_widths[-1]) * total_chars))
            text = _stringify_cell(value)
            out_row.append(textwrap.fill(text, width=width, break_long_words=True, break_on_hyphens=True))
        wrapped.append(out_row)
    return wrapped


def _row_line_units(row: list[str]) -> int:
    lines = [str(cell).count("\n") + 1 for cell in row]
    return max(lines) if lines else 1


def _chunk_table_rows_for_pdf(rows: list[list[str]], *, line_budget: int = 30) -> list[list[list[str]]]:
    if not rows:
        return []
    chunks: list[list[list[str]]] = []
    current: list[list[str]] = []
    used = 0
    for row in rows:
        units = _row_line_units(row)
        if current and (used + units) > line_budget:
            chunks.append(current)
            current = []
            used = 0
        current.append(row)
        used += units
    if current:
        chunks.append(current)
    return chunks


def _wrap_line(line: str, width: int = 95) -> list[str]:
    text = str(line)
    if not text:
        return [""]
    return textwrap.wrap(text, width=width) or [text]


def _section_lines(section: dict[str, Any]) -> list[str]:
    lines: list[str] = []
    title = str(section.get("title") or "").strip()
    if title:
        lines.append(title)
        lines.append("-" * len(title))

    key_values = section.get("key_values")
    if isinstance(key_values, dict):
        for key, value in key_values.items():
            lines.extend(_wrap_line(f"{key}: {value}"))

    paragraphs = section.get("paragraphs")
    if isinstance(paragraphs, list):
        for item in paragraphs:
            lines.extend(_wrap_line(str(item)))
            lines.append("")

    bullets = section.get("bullets")
    if isinstance(bullets, list):
        for item in bullets:
            lines.extend(_wrap_line(f"- {item}"))

    if "json_block" in section:
        lines.append("")
        lines.append("JSON")
        lines.append("----")
        lines.extend(_to_json_text(section.get("json_block")).splitlines())

    if lines and lines[-1] != "":
        lines.append("")
    return lines


def _write_pdf(path: Path, payload: dict[str, Any]) -> None:
    title = str(payload.get("title") or "Analysis Report")
    subtitle = str(payload.get("subtitle") or "")
    sections = payload.get("sections") if isinstance(payload.get("sections"), list) else []
    figures = payload.get("figures") if isinstance(payload.get("figures"), list) else []

    with PdfPages(path) as pdf:
        first_page = True
        if not sections:
            fig = plt.figure(figsize=(8.27, 11.69))
            ax = fig.add_axes([0.0, 0.0, 1.0, 1.0])
            ax.set_axis_off()
            ax.text(0.05, 0.97, title, fontsize=16, fontweight="bold", va="top")
            if subtitle:
                ax.text(0.05, 0.93, subtitle, fontsize=10, va="top")
            ax.text(0.05, 0.88, "No report content.", fontsize=10, va="top")
            pdf.savefig(fig)
            plt.close(fig)
            first_page = False

        for section in sections:
            if not isinstance(section, dict):
                continue
            section_title = str(section.get("title") or "").strip()
            table_data = _section_table_data(section)

            if table_data is not None:
                headers, rows = table_data
                if not rows:
                    rows = [["No data"] + [""] * max(0, len(headers) - 1)]
                col_widths = _table_col_widths(len(headers))
                wrapped_rows = _wrap_table_rows(rows, col_widths)
                row_chunks = _chunk_table_rows_for_pdf(wrapped_rows, line_budget=30)
                for chunk in row_chunks:
                    fig = plt.figure(figsize=(8.27, 11.69))
                    ax = fig.add_axes([0.0, 0.0, 1.0, 1.0])
                    ax.set_axis_off()

                    y = 0.97
                    ax.text(0.05, y, title, fontsize=14 if first_page else 12, fontweight="bold", va="top")
                    y -= 0.04
                    if first_page and subtitle:
                        ax.text(0.05, y, subtitle, fontsize=9.5, va="top")
                        y -= 0.035
                    if section_title:
                        ax.text(0.05, y, section_title, fontsize=11, fontweight="bold", va="top")

                    t_ax = fig.add_axes([0.05, 0.08, 0.90, 0.78])
                    t_ax.set_axis_off()
                    tbl = t_ax.table(
                        cellText=chunk,
                        colLabels=headers,
                        colWidths=col_widths,
                        cellLoc="left",
                        colLoc="left",
                        loc="upper center",
                    )
                    tbl.auto_set_font_size(False)
                    tbl.set_fontsize(7.8)
                    tbl.scale(1.0, 1.15)
                    for (r, _c), cell in tbl.get_celld().items():
                        cell.set_edgecolor("black")
                        cell.set_linewidth(0.8)
                        cell.get_text().set_wrap(True)
                        if r == 0:
                            cell.set_text_props(fontweight="bold")
                    # Increase row heights according to wrapped line count.
                    base_h = 0.028
                    for ridx, row_values in enumerate(chunk, start=1):
                        h = base_h * float(max(1, _row_line_units(row_values)))
                        for cidx in range(len(headers)):
                            if (ridx, cidx) in tbl.get_celld():
                                tbl[(ridx, cidx)].set_height(h)
                    for cidx in range(len(headers)):
                        if (0, cidx) in tbl.get_celld():
                            tbl[(0, cidx)].set_height(base_h * 1.1)

                    pdf.savefig(fig)
                    plt.close(fig)
                    first_page = False

            # Render non-table narrative block (bullets/paragraphs/json_block) if any.
            text_section = dict(section)
            text_section.pop("key_values", None)
            text_section.pop("table", None)
            if any(
                text_section.get(k)
                for k in ("paragraphs", "bullets")
            ) or ("json_block" in text_section):
                lines = _section_lines(text_section)
                max_lines_per_page = 50
                chunks = [
                    lines[i : i + max_lines_per_page]
                    for i in range(0, len(lines), max_lines_per_page)
                ]
                for chunk in chunks:
                    fig = plt.figure(figsize=(8.27, 11.69))
                    ax = fig.add_axes([0.0, 0.0, 1.0, 1.0])
                    ax.set_axis_off()

                    y = 0.97
                    ax.text(0.05, y, title, fontsize=14 if first_page else 12, fontweight="bold", va="top")
                    y -= 0.04
                    if first_page and subtitle:
                        ax.text(0.05, y, subtitle, fontsize=9.5, va="top")
                        y -= 0.035
                    for line in chunk:
                        ax.text(0.05, y, line, fontsize=9.5, va="top")
                        y -= 0.017

                    pdf.savefig(fig)
                    plt.close(fig)
                    first_page = False

        for item in figures:
            if not isinstance(item, dict):
                continue
            fig_path_raw = item.get("path")
            fig_path = Path(str(fig_path_raw)) if fig_path_raw else None
            if fig_path is None or not fig_path.exists():
                continue
            caption = str(item.get("caption") or fig_path.name)
            try:
                image = plt.imread(fig_path)
            except Exception:
                continue
            fig = plt.figure(figsize=(8.27, 11.69))
            ax = fig.add_axes([0.05, 0.12, 0.90, 0.80])
            ax.set_axis_off()
            ax.imshow(image)
            fig.text(0.05, 0.05, caption, fontsize=10)
            pdf.savefig(fig)
            plt.close(fig)


def _write_docx(path: Path, payload: dict[str, Any]) -> None:
    from docx import Document  # type: ignore
    from docx.shared import Inches  # type: ignore

    title = str(payload.get("title") or "Analysis Report")
    subtitle = str(payload.get("subtitle") or "")
    sections = payload.get("sections") if isinstance(payload.get("sections"), list) else []
    figures = payload.get("figures") if isinstance(payload.get("figures"), list) else []

    doc = Document()
    doc.add_heading(title, level=0)
    if subtitle:
        doc.add_paragraph(subtitle)

    for section in sections:
        if not isinstance(section, dict):
            continue
        sec_title = str(section.get("title") or "").strip()
        if sec_title:
            doc.add_heading(sec_title, level=1)

        table_data = _section_table_data(section)
        if table_data is not None:
            headers, rows = table_data
            col_widths = _table_col_widths(len(headers))
            wrapped_rows = _wrap_table_rows(rows, col_widths)
            table = doc.add_table(rows=1, cols=max(1, len(headers)))
            table.style = "Table Grid"
            table.autofit = False
            for i, header in enumerate(headers):
                table.rows[0].cells[i].text = str(header)
                table.rows[0].cells[i].width = Inches(6.5 * col_widths[i if i < len(col_widths) else -1])
            for values in wrapped_rows:
                row = table.add_row().cells
                for i, value in enumerate(values[: len(headers)]):
                    row[i].text = str(value)
                    row[i].width = Inches(6.5 * col_widths[i if i < len(col_widths) else -1])

        paragraphs = section.get("paragraphs")
        if isinstance(paragraphs, list):
            for item in paragraphs:
                doc.add_paragraph(str(item))

        bullets = section.get("bullets")
        if isinstance(bullets, list):
            for item in bullets:
                doc.add_paragraph(str(item), style="List Bullet")

        if "json_block" in section:
            doc.add_paragraph(_to_json_text(section.get("json_block")))

    if figures:
        doc.add_heading("Figures", level=1)
    for item in figures:
        if not isinstance(item, dict):
            continue
        fig_path_raw = item.get("path")
        fig_path = Path(str(fig_path_raw)) if fig_path_raw else None
        if fig_path is None or not fig_path.exists():
            continue
        caption = str(item.get("caption") or fig_path.name)
        doc.add_paragraph(caption)
        doc.add_picture(str(fig_path), width=Inches(6.5))

    doc.save(path)


def write_report_artifacts(
    payload: dict[str, Any],
    *,
    out_dir: Path,
    stem: str,
    formats: Sequence[str],
) -> tuple[list[str], list[str]]:
    """Write report artifacts and return (written_files, notes)."""
    out_dir.mkdir(parents=True, exist_ok=True)
    normalized = _normalize_payload(payload)

    written: list[str] = []
    notes: list[str] = []

    json_path = out_dir / f"{stem}_report_data.json"
    json_path.write_text(_to_json_text(normalized), encoding="utf-8")
    written.append(json_path.name)

    targets = tuple(str(f).lower() for f in formats)
    if "pdf" in targets:
        pdf_path = out_dir / f"{stem}_report.pdf"
        _write_pdf(pdf_path, normalized)
        written.append(pdf_path.name)

    if "docx" in targets:
        docx_path = out_dir / f"{stem}_report.docx"
        try:
            _write_docx(docx_path, normalized)
            written.append(docx_path.name)
        except Exception as exc:
            notes.append(f"DOCX report not written: {exc}")

    return written, notes


__all__ = [
    "normalize_report_formats",
    "write_report_artifacts",
]
